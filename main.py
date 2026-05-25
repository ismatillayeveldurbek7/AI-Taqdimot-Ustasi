import os
import csv
import re
import asyncio
import logging
import sqlite3
import zipfile
from datetime import datetime, timezone, timedelta
from html import escape
from typing import Optional

from aiogram import Bot, Dispatcher, F
from aiogram.enums import ChatMemberStatus
from aiogram.filters import Command
from aiogram.exceptions import TelegramBadRequest
from aiogram.types import (
    Message,
    CallbackQuery,
    InlineKeyboardMarkup,
    InlineKeyboardButton,
    FSInputFile,
)
from aiogram.utils.keyboard import InlineKeyboardBuilder

try:
    from openpyxl import Workbook
except ImportError:
    Workbook = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    Document = None

# =========================
# SOZLAMALAR
# =========================
# FIX #1: Token faqat env dan olinadi, kodga yozilmaydi
BOT_TOKEN = os.getenv("BOT_TOKEN")
if not BOT_TOKEN:
    raise ValueError("BOT_TOKEN muhit o'zgaruvchisi topilmadi. Iltimos, BOT_TOKEN ni o'rnating.")

CHANNEL_USERNAME = os.getenv("CHANNEL_USERNAME", "@Qashqadaryo_PMM")
CHANNEL_URL = "https://t.me/Qashqadaryo_PMM"

ADMIN_IDS_RAW = os.getenv("ADMIN_IDS", "5298063089,7361393654")
ADMIN_IDS = [int(x.strip()) for x in ADMIN_IDS_RAW.split(",") if x.strip().isdigit()]

FACEBOOK_URL = "https://www.facebook.com/share/1E4ZVePTh4/"
INSTAGRAM_URL = "https://www.instagram.com/pedagogikmahorat"

DATA_DIR = os.getenv("DATA_DIR", "/app/data")
os.makedirs(DATA_DIR, exist_ok=True)

DB_NAME = os.path.join(DATA_DIR, "votes.db")
EXPORT_FILE = os.path.join(DATA_DIR, "votes_export.csv")
VOTES_XLSX_FILE = os.path.join(DATA_DIR, "votes_export.xlsx")
RATING_XLSX_FILE = os.path.join(DATA_DIR, "rating_export.xlsx")
# FIX #8: Ikki xil nom o'rniga bitta nom
COMPLAINTS_DOCX_FILE = os.path.join(DATA_DIR, "complaints_export.docx")
BACKUP_ZIP_FILE = os.path.join(DATA_DIR, "bot_backup.zip")

SUBJECTS = {
    "s1": {
        "name": "Tillarni o'qitish metodikasi",
        "old_key": "tillarni_oqitish_metodikasi",
        "teachers": {
            "tom_1": "Norov Otajon Shomurodovich",
            "tom_2": "Abdixolikov Abdulazizxon Abduvohob o'g'li",
            "tom_3": "Azimova Nigora Anvar qizi",
            "tom_4": "Abatov Doston Ro'zimurod o'g'li",
            "tom_5": "Jalilova Komila Abdullayevna",
            "tom_6": "Oqboyeva Zulfiya Bobonazarovna",
            "tom_7": "Sevastyanova Nadejda Aleksandrovna",
            "tom_8": "Xidirova Feruza To'rayevna",
            "tom_9": "Ergasheva Dilorom Muradilloyevna",
        },
    },
    "s2": {
        "name": "Pedagogika, psixologiya va ta'lim menejmenti",
        "old_key": "pedagogika_psixologiya_va_talim_menejmenti",
        "teachers": {
            "pptm_1": "Umarov Lutfillo Murodilloyevich",
            "pptm_2": "Baratova Nasiba Turobovna",
            "pptm_3": "Bekmurodova Dilnoza Pirimovna",
            "pptm_4": "Meyliyev Lobar Nurmatovna",
            "pptm_5": "Ochilov Og'abek Narzullayevich",
            "pptm_6": "Shoniyozova Dilafruz Sabirovna",
            "pptm_7": "Yaratov Xamidjon Muxtorovich",
            "pptm_8": "Nazarov Asliddin Faxriddin o'g'li",
            "pptm_9": "Ergasheva Dilafruz Ergamqulovna",
            "pptm_10": "Soatov Asadulloh Jabborovich",
        },
    },
    "s3": {
        "name": "Aniq va tabiiy fanlar",
        "old_key": "aniq_va_tabiiy_fanlar",
        "teachers": {
            "atf_1": "Jobborov Farhod Bo'rinevich",
            "atf_2": "Karimova Habiba Abduraxmonovna",
            "atf_3": "Quldoshova Maftuna Jumanzar qizi",
            "atf_4": "Mallaev Xamro Ro'ziboyevich",
            "atf_5": "Mamatov Bekzod Farxotovich",
            "atf_6": "Pardaeva Muqaddas Zafar qizi",
            "atf_7": "Parmanov Jahongir Rayhonovich",
            "atf_8": "Rahmatullayev Erkin Shokirovich",
            "atf_9": "Suyarov Zoir Shojmardonovich",
            "atf_10": "Tursunova Maftuna Sulton qizi",
            "atf_11": "Umarov Ibrohimxon Norxuja o'g'li",
            "atf_12": "Chariev Rashid Ravshanovich",
            "atf_13": "Elmurodov Sherdil Ergashyevich",
            "atf_14": "Eshmonov Laziz Norxo'rja o'g'li",
            "atf_15": "Karaeva Dilfuzaxon Mamasharipovna",
            "atf_16": "Salomova Madina Sodiq qizi",
        },
    },
    "s4": {
        "name": "Amaliy va ijtimoiy fanlar",
        "old_key": "amaliy_va_ijtimoiy_fanlar",
        "teachers": {
            "aif_1": "Yo'ldashev Bekmirza Elmurodovich",
            "aif_2": "Jabboborov Laziz Hamza o'g'li",
            "aif_3": "Nurmatov Samandar Fayratovich",
            "aif_4": "Batoshov Inatillo Kungirovich",
            "aif_5": "Rajabov Ruslan Bozorovich",
            "aif_6": "Sanaev Azamat Alponovich",
            "aif_7": "Shamsiev Jahongir Qulmurod o'g'li",
            "aif_8": "Xudoyberdiev Axrorboy Nabi o'g'li",
            "aif_9": "Xasanova Gulnora Qorshanbiyevna",
            "aif_10": "Eshnazarova Maziya Allanazarovna",
        },
    },
    "s5": {
        "name": "Maktabgacha, boshlang'ich va maxsus ta'lim",
        "old_key": "maktabgacha_boshlangich_va_maxsus_talim",
        "teachers": {
            "mbmt_1": "Irisova Sayyora Rajabovna",
            "mbmt_2": "Azizova Dilnoz Yo'ldoshevna",
            "mbmt_3": "G'oyimov Umar Eshmurodovich",
            "mbmt_4": "Ziyotova Madina Mansur qizi",
            "mbmt_5": "Karimova Umida Sharopovna",
            "mbmt_6": "Qarshiyeva Guzal Alimardonovna",
            "mbmt_7": "Qurbanova Xusnora Xudoyberdi qizi",
            "mbmt_8": "Rajabova Xurshida Hakimovna",
            "mbmt_9": "Razzaqova Dilnoza Akramovna",
            "mbmt_10": "Sadinova Marjona Akmal qizi",
            "mbmt_11": "Shaxmurodova Dilxaxon Almardanovna",
            "mbmt_12": "Ergasheva Xusniya Mirzoxid qizi",
            "mbmt_13": "Zaripova Muslima Qurbonovna",
        },
    },
}
OLD_TO_NEW_SUBJECT = {v["old_key"]: k for k, v in SUBJECTS.items()}

logging.basicConfig(level=logging.INFO)

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

conn = sqlite3.connect(DB_NAME, check_same_thread=False)
cursor = conn.cursor()
db_lock = asyncio.Lock()
WAITING_COMPLAINT_TEXT = set()
COMPLAINT_COOLDOWN_SECONDS = 300
COMPLAINT_MAX_LENGTH = 1000

# Refresh tugmalarini ketma-ket bosishdan himoya
LAST_REFRESH = {}
REFRESH_BUSY = set()
REFRESH_COOLDOWN_SECONDS = 1.5


# =========================
# O'ZBEKISTON VAQTI
# =========================
UZ_TZ = timezone(timedelta(hours=5), name="Asia/Tashkent")

def uz_now() -> datetime:
    return datetime.now(UZ_TZ)

# =========================
# LOTIN / KRILL
# =========================
def latin_to_cyrillic_text(text: str) -> str:
    # FIX #4: Takrorlangan juftlar olib tashlandi, to'g'ri tartib saqlandi
    pairs = [
        ("O\u2018", "\u040e"), ("o\u2018", "\u045e"),   # O' → Ў (unicode apostrof)
        ("G\u2018", "\u0492"), ("g\u2018", "\u0493"),   # G' → Ғ
        ("O'", "\u040e"), ("o'", "\u045e"),              # O' → Ў (oddiy apostrof)
        ("G'", "\u0492"), ("g'", "\u0493"),              # G' → Ғ
        ("Sh", "\u0428"), ("sh", "\u0448"),
        ("Ch", "\u0427"), ("ch", "\u0447"),
        ("Ya", "\u042f"), ("ya", "\u044f"),
        ("Yo", "\u0401"), ("yo", "\u0451"),
        ("Yu", "\u042e"), ("yu", "\u044e"),
        ("Ts", "\u0426"), ("ts", "\u0446"),
    ]
    for old, new in pairs:
        text = text.replace(old, new)

    table = str.maketrans({
        "A": "\u0410", "a": "\u0430",
        "B": "\u0411", "b": "\u0431",
        "D": "\u0414", "d": "\u0434",
        "E": "\u0415", "e": "\u0435",
        "F": "\u0424", "f": "\u0444",
        "G": "\u0413", "g": "\u0433",
        "H": "\u04b2", "h": "\u04b3",
        "I": "\u0418", "i": "\u0438",
        "J": "\u0416", "j": "\u0436",
        "K": "\u041a", "k": "\u043a",
        "L": "\u041b", "l": "\u043b",
        "M": "\u041c", "m": "\u043c",
        "N": "\u041d", "n": "\u043d",
        "O": "\u041e", "o": "\u043e",
        "P": "\u041f", "p": "\u043f",
        "Q": "\u049a", "q": "\u049b",
        "R": "\u0420", "r": "\u0440",
        "S": "\u0421", "s": "\u0441",
        "T": "\u0422", "t": "\u0442",
        "U": "\u0423", "u": "\u0443",
        "V": "\u0412", "v": "\u0432",
        "X": "\u0425", "x": "\u0445",
        "Y": "\u0419", "y": "\u0439",
        "Z": "\u0417", "z": "\u0437",
        "`": "\u044a", "'": "\u044a", "\u2019": "\u044a",
    })
    return text.translate(table)


def translit_html_safe(text: str, script: str) -> str:
    parts = re.split(r"(<[^>]+>)", text)
    result = []
    for part in parts:
        if part.startswith("<") and part.endswith(">"):
            result.append(part)
        else:
            result.append(latin_to_cyrillic_text(part) if script == "cyrillic" else part)
    return "".join(result)


def get_user_script(user_id: int) -> str:
    ensure_user(user_id)
    cursor.execute("SELECT script FROM user_prefs WHERE user_id = ?", (user_id,))
    row = cursor.fetchone()
    return row[0] if row and row[0] in ("latin", "cyrillic") else "latin"


def set_user_script(user_id: int, script: str):
    ensure_user(user_id)
    if script not in ("latin", "cyrillic"):
        script = "latin"
    cursor.execute("UPDATE user_prefs SET script = ? WHERE user_id = ?", (script, user_id))
    conn.commit()


def tr(user_id: int, text: str) -> str:
    return translit_html_safe(text, get_user_script(user_id))

def like_label(user_id: int) -> str:
    return "Лайк" if get_user_script(user_id) == "cyrillic" else "Like"

def dislike_label(user_id: int) -> str:
    return "Дислайк" if get_user_script(user_id) == "cyrillic" else "Dislike"


# =========================
# DB
# =========================
def normalize_subject_key(subject_key: str) -> str:
    if subject_key == "general":
        return "general"
    return OLD_TO_NEW_SUBJECT.get(subject_key, subject_key)

def init_db():
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS votes (
            user_id INTEGER PRIMARY KEY,
            full_name TEXT,
            username TEXT,
            subject_key TEXT NOT NULL,
            teacher_key TEXT NOT NULL,
            voted_at TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS user_prefs (
            user_id INTEGER PRIMARY KEY,
            script TEXT DEFAULT 'latin',
            access_granted INTEGER DEFAULT 0
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS teacher_ratings (
            user_id INTEGER NOT NULL,
            full_name TEXT,
            username TEXT,
            subject_key TEXT NOT NULL,
            teacher_key TEXT NOT NULL,
            rating TEXT NOT NULL CHECK(rating IN ('like', 'dislike')),
            rated_at TEXT,
            PRIMARY KEY (user_id, subject_key, teacher_key)
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS complaints (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            full_name TEXT,
            username TEXT,
            message_text TEXT NOT NULL,
            created_at TEXT
        )
    """)
    conn.commit()
    migrate_old_subject_keys()
    if get_setting("voting_open", "") == "":
        set_setting("voting_open", "1")

def migrate_old_subject_keys():
    for old_key, new_key in OLD_TO_NEW_SUBJECT.items():
        cursor.execute("UPDATE votes SET subject_key = ? WHERE subject_key = ?", (new_key, old_key))
        cursor.execute("UPDATE teacher_ratings SET subject_key = ? WHERE subject_key = ?", (new_key, old_key))
    conn.commit()

def get_setting(key: str, default: str = "") -> str:
    cursor.execute("SELECT value FROM settings WHERE key = ?", (key,))
    row = cursor.fetchone()
    return row[0] if row else default

def set_setting(key: str, value: str):
    cursor.execute("""
        INSERT INTO settings (key, value)
        VALUES (?, ?)
        ON CONFLICT(key) DO UPDATE SET value = excluded.value
    """, (key, value))
    conn.commit()

def ensure_user(user_id: int):
    cursor.execute("""
        INSERT INTO user_prefs (user_id, script, access_granted)
        VALUES (?, 'latin', 0)
        ON CONFLICT(user_id) DO NOTHING
    """, (user_id,))
    conn.commit()

def has_access(user_id: int) -> bool:
    ensure_user(user_id)
    cursor.execute("SELECT access_granted FROM user_prefs WHERE user_id = ?", (user_id,))
    row = cursor.fetchone()
    return bool(row[0]) if row else False

def require_access_only(user_id: int) -> bool:
    return has_access(user_id)

def grant_access(user_id: int):
    ensure_user(user_id)
    cursor.execute("UPDATE user_prefs SET access_granted = 1 WHERE user_id = ?", (user_id,))
    conn.commit()

def reset_access(user_id: int):
    ensure_user(user_id)
    cursor.execute("UPDATE user_prefs SET access_granted = 0 WHERE user_id = ?", (user_id,))
    conn.commit()

def is_admin(user_id: int) -> bool:
    return user_id in ADMIN_IDS

def is_voting_open() -> bool:
    return get_setting("voting_open", "1") == "1"

def open_voting():
    set_setting("voting_open", "1")

def close_voting():
    set_setting("voting_open", "0")

def has_voted(user_id: int) -> bool:
    cursor.execute("SELECT 1 FROM votes WHERE user_id = ?", (user_id,))
    return cursor.fetchone() is not None

def save_vote(user_id: int, full_name: str, username: str, subject_key: str, teacher_key: str) -> bool:
    """
    Ovoz saqlaydi.
    True  = ovoz muvaffaqiyatli saqlandi.
    False = foydalanuvchi oldin ovoz bergan yoki PRIMARY KEY to'qnashuvi bo'lgan.
    """
    subject_key = normalize_subject_key(subject_key)

    try:
        cursor.execute("""
            INSERT INTO votes (user_id, full_name, username, subject_key, teacher_key, voted_at)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (
            user_id,
            full_name,
            username,
            subject_key,
            teacher_key,
            uz_now().strftime("%Y-%m-%d %H:%M:%S")
        ))
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False

# FIX #9: str | None o'rniga Optional[str] — Python 3.9 bilan moslik
def get_total_votes(subject_key: Optional[str] = None) -> int:
    if subject_key:
        cursor.execute("SELECT COUNT(*) FROM votes WHERE subject_key = ?", (normalize_subject_key(subject_key),))
    else:
        cursor.execute("SELECT COUNT(*) FROM votes")
    return cursor.fetchone()[0]

def reset_votes():
    cursor.execute("DELETE FROM votes")
    conn.commit()

def reset_ratings():
    cursor.execute("DELETE FROM teacher_ratings")
    conn.commit()

def reset_complaints():
    cursor.execute("DELETE FROM complaints")
    conn.commit()

def get_subject_name(subject_key: str) -> str:
    subject_key = normalize_subject_key(subject_key)
    return SUBJECTS.get(subject_key, {}).get("name", subject_key)

def get_teacher_name(subject_key: str, teacher_key: str) -> str:
    subject_key = normalize_subject_key(subject_key)
    return SUBJECTS.get(subject_key, {}).get("teachers", {}).get(teacher_key, teacher_key)

def build_progress_bar(percent: float, length: int = 14) -> str:
    filled = round((percent / 100) * length)
    filled = max(0, min(filled, length))
    return "\u2593" * filled + "\u2591" * (length - filled)

def get_all_teachers_flat():
    items = []
    for subject_key, subject_data in SUBJECTS.items():
        for teacher_key, teacher_name in subject_data["teachers"].items():
            items.append((subject_key, teacher_key, teacher_name))
    return items

# =========================
# RATING DB / STATS
# =========================
def save_teacher_rating(user_id: int, full_name: str, username: str, subject_key: str, teacher_key: str, rating: str):
    subject_key = normalize_subject_key(subject_key)
    cursor.execute("""
        INSERT INTO teacher_ratings (user_id, full_name, username, subject_key, teacher_key, rating, rated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(user_id, subject_key, teacher_key)
        DO UPDATE SET
            full_name = excluded.full_name,
            username = excluded.username,
            rating = excluded.rating,
            rated_at = excluded.rated_at
    """, (user_id, full_name, username, subject_key, teacher_key, rating, uz_now().strftime("%Y-%m-%d %H:%M:%S")))
    conn.commit()

# FIX #9: Optional ishlatildi
def get_user_teacher_rating(user_id: int, subject_key: str, teacher_key: str) -> Optional[str]:
    cursor.execute("""
        SELECT rating FROM teacher_ratings
        WHERE user_id = ? AND subject_key = ? AND teacher_key = ?
    """, (user_id, normalize_subject_key(subject_key), teacher_key))
    row = cursor.fetchone()
    return row[0] if row else None

def get_rating_counts(subject_key: str, teacher_key: str):
    cursor.execute("""
        SELECT
            SUM(CASE WHEN rating = 'like' THEN 1 ELSE 0 END),
            SUM(CASE WHEN rating = 'dislike' THEN 1 ELSE 0 END),
            COUNT(*)
        FROM teacher_ratings
        WHERE subject_key = ? AND teacher_key = ?
    """, (normalize_subject_key(subject_key), teacher_key))
    like_count, dislike_count, total = cursor.fetchone()
    like_count = like_count or 0
    dislike_count = dislike_count or 0
    total = total or 0
    like_percent = (like_count / total * 100) if total else 0
    dislike_percent = (dislike_count / total * 100) if total else 0
    return like_count, dislike_count, total, like_percent, dislike_percent

def rating_rows():
    rows = []
    for subject_key, teacher_key, teacher_name in get_all_teachers_flat():
        like_count, dislike_count, total, like_percent, dislike_percent = get_rating_counts(subject_key, teacher_key)
        rows.append({
            "subject_key": subject_key,
            "subject_name": get_subject_name(subject_key),
            "teacher_key": teacher_key,
            "teacher_name": teacher_name,
            "like": like_count,
            "dislike": dislike_count,
            "total": total,
            "like_percent": like_percent,
            "dislike_percent": dislike_percent,
        })
    return rows

def get_vote_percent(count: int, denominator: int) -> float:
    return (count / denominator * 100) if denominator > 0 else 0.0

def save_complaint(user_id: int, full_name: str, username: str, message_text: str):
    cursor.execute("""
        INSERT INTO complaints (user_id, full_name, username, message_text, created_at)
        VALUES (?, ?, ?, ?, ?)
    """, (
        user_id,
        full_name,
        username,
        message_text,
        uz_now().strftime("%Y-%m-%d %H:%M:%S")
    ))
    conn.commit()


def get_last_complaint_for_user(user_id: int):
    cursor.execute("""
        SELECT message_text, created_at
        FROM complaints
        WHERE user_id = ?
        ORDER BY id DESC
        LIMIT 1
    """, (user_id,))
    return cursor.fetchone()


def complaint_allowed(user_id: int, message_text: str):
    """
    Shikoyat/taklif spam himoyasi.
    True, '' qaytarilsa yuborish mumkin. Aks holda False va xabar matni qaytadi.
    """
    text = (message_text or "").strip()
    if len(text) > COMPLAINT_MAX_LENGTH:
        return False, f"Xabar juda uzun. Iltimos, {COMPLAINT_MAX_LENGTH} ta belgidan oshirmang."

    last = get_last_complaint_for_user(user_id)
    if not last:
        return True, ""

    last_text, created_at = last
    if (last_text or "").strip() == text:
        return False, "Siz aynan shu xabarni avval yuborgansiz. Iltimos, takroriy xabar yubormang."

    try:
        last_dt = datetime.strptime(created_at or "", "%Y-%m-%d %H:%M:%S").replace(tzinfo=UZ_TZ)
        diff = (uz_now() - last_dt).total_seconds()
        if diff < COMPLAINT_COOLDOWN_SECONDS:
            wait_seconds = int(COMPLAINT_COOLDOWN_SECONDS - diff)
            minutes = max(1, (wait_seconds + 59) // 60)
            return False, f"Spamdan himoya uchun keyingi murojaatni taxminan {minutes} daqiqadan keyin yuboring."
    except Exception:
        pass

    return True, ""

# FIX #9: Optional ishlatildi
def get_complaints_rows(limit: Optional[int] = None):
    sql = """
        SELECT id, user_id, full_name, username, message_text, created_at
        FROM complaints
        ORDER BY id DESC
    """
    params = ()
    if limit:
        sql += " LIMIT ?"
        params = (limit,)
    cursor.execute(sql, params)
    return cursor.fetchall()


def get_complaints_count() -> int:
    cursor.execute("SELECT COUNT(*) FROM complaints")
    return cursor.fetchone()[0]


def get_complaints_text(user_id: int) -> str:
    rows = get_complaints_rows(limit=30)
    total = get_complaints_count()

    if not rows:
        return tr(user_id, "📩 <b>Shikoyat va takliflar</b>\n\nHali hech qanday xabar kelmagan.")

    lines = [f"📩 <b>Shikoyat va takliflar</b>\n\nJami: {total} ta\nOxirgi {len(rows)} ta xabar:\n"]
    for i, (cid, uid, full_name, username, message_text, created_at) in enumerate(rows, start=1):
        safe_name = escape(full_name or "Noma'lum")
        safe_username = escape(username or "")
        safe_message = escape(message_text or "")
        line = f"{i}. <b>{safe_name}</b>"
        if safe_username:
            line += f" (@{safe_username})"
        line += f"\n   ID: <code>{uid}</code>"
        line += f"\n   Sana: {escape(created_at or '')}"
        line += f"\n   Xabar: {safe_message}"
        lines.append(line)

    text = "\n\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n\n... qisqartirildi"
    return tr(user_id, text)


def export_complaints_to_docx() -> str:
    if Document is None:
        path = os.path.join(DATA_DIR, "complaints_export.txt")
        rows = get_complaints_rows()
        with open(path, "w", encoding="utf-8") as f:
            f.write("Shikoyat va takliflar\n")
            f.write(f"Jami: {len(rows)} ta\n\n")
            for i, (cid, uid, full_name, username, message_text, created_at) in enumerate(rows, start=1):
                # FIX #3: f-string ichida apostrof — o'zgaruvchiga olindi
                name = full_name or "Noma'lum"
                uname = f"@{username}" if username else ""
                f.write(f"{i}. {name} ({uname})\n")
                f.write(f"ID: {uid}\nSana: {created_at or ''}\nXabar: {message_text or ''}\n\n")
        return path

    rows = get_complaints_rows()
    doc = Document()

    title = doc.add_heading("Shikoyat va takliflar", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Jami xabarlar soni: ").bold = True
    p.add_run(str(len(rows)))

    p = doc.add_paragraph()
    p.add_run("Yaratilgan sana: ").bold = True
    p.add_run(uz_now().strftime("%Y-%m-%d %H:%M:%S"))

    if not rows:
        doc.add_paragraph("Hali hech qanday shikoyat yoki taklif kelmagan.")
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["№", "F.I.Sh", "Username", "Telegram ID", "Sana", "Xabar"]
        for idx, header in enumerate(headers):
            run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
            run.bold = True

        for i, (cid, uid, full_name, username, message_text, created_at) in enumerate(rows, start=1):
            cells = table.add_row().cells
            cells[0].text = str(i)
            # FIX #3: apostrof xatosi bartaraf etildi
            cells[1].text = full_name or "Noma'lum"
            cells[2].text = f"@{username}" if username else ""
            cells[3].text = str(uid)
            cells[4].text = created_at or ""
            cells[5].text = message_text or ""

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)

    doc.save(COMPLAINTS_DOCX_FILE)
    return COMPLAINTS_DOCX_FILE

def get_subscription_required_alert(user_id: int) -> str:
    return tr(user_id, "Avval Telegram kanalga obuna bo'ling va ✅ Tekshirish tugmasini bosing.")

# =========================
# TEXTS
# =========================
def get_welcome_text(user_id: int) -> str:
    return tr(
        user_id,
        "🚀 <b>Botdan foydalanish uchun quyidagilarni bajaring:</b>\n\n"
        "1️⃣ 📸 Instagram sahifaga o'ting\n"
        "2️⃣ 📘 Facebook sahifaga o'ting\n"
        "3️⃣ 📢 Telegram kanalga obuna bo'ling\n\n"
        "Telegram obunasi bot tomonidan tekshiriladi. Instagram/Facebook tugmalari URL sifatida berilgan.\n\n"
        "👇 Telegram kanalga obuna bo'lgach, <b>✅ Tekshirish</b> tugmasini bosing."
    )

def get_home_text(user_id: int) -> str:
    return tr(user_id, "🏠 <b>Bosh menyu</b>\n\nKerakli bo'limni tanlang:")

def get_help_text(user_id: int) -> str:
    return tr(
        user_id,
        "ℹ️ <b>Yordam</b>\n\n"
        "• Telegram kanalga obuna bo'ling\n"
        "• So'ng ✅ Tekshirish tugmasini bosing\n"
        "• Ovoz berish uchun kafedra tanlanadi\n"
        "• Keyin o'qituvchi tanlanadi\n"
        "• Har bir foydalanuvchi asosiy ovozni faqat 1 marta beradi\n"
        "• O'qituvchilarni baholashda har bir o'qituvchiga 1 ta like/dislike beriladi\n"
        "• Bahoni keyin o'zgartirish mumkin\n"
        "• Natijalarni istalgan payt ko'rishingiz mumkin"
    )

def get_already_voted_text(user_id: int) -> str:
    return tr(user_id, "✅ <b>Siz allaqachon ovoz berib bo'lgansiz</b>\n\nQayta ovoz berish mumkin emas.")

def get_closed_text(user_id: int) -> str:
    return tr(user_id, "🔒 <b>Ovoz berish hozircha yopilgan</b>\n\nAdmin tomonidan ovoz berish vaqtincha to'xtatilgan.")

def get_subject_select_text(user_id: int) -> str:
    return tr(user_id, "🗂 <b>Kafedrani tanlang</b>\n\nQuyidagi bo'limlardan birini tanlang:")

def get_teacher_select_text(user_id: int, subject_key: str) -> str:
    return tr(user_id, f"{SUBJECTS[subject_key]['name']}\n\n<b>O'qituvchini tanlang:</b>")

def get_rating_select_text(user_id: int) -> str:
    return tr(user_id, "⭐️ <b>O'qituvchilarni baholash</b>\n\nAvval kafedrani tanlang:")

def get_rating_teacher_text(user_id: int, subject_key: str) -> str:
    return tr(user_id, f"⭐️ <b>{SUBJECTS[subject_key]['name']}</b>\n\nBaholash uchun o'qituvchini tanlang:")

def get_rate_text(user_id: int, subject_key: str, teacher_key: str) -> str:
    current = get_user_teacher_rating(user_id, subject_key, teacher_key)
    if current == "like":
        current_text = f"Hozirgi bahoyingiz: 👍 {like_label(user_id)}"
    elif current == "dislike":
        current_text = f"Hozirgi bahoyingiz: 👎 {dislike_label(user_id)}"
    else:
        current_text = "Hozirgi bahoyingiz: hali baho berilmagan"
    like_count, dislike_count, total, like_percent, dislike_percent = get_rating_counts(subject_key, teacher_key)
    return tr(
        user_id,
        f"⭐️ <b>O'qituvchini baholash</b>\n\n"
        f"<b>Kafedra:</b> {get_subject_name(subject_key)}\n"
        f"<b>O'qituvchi:</b> {get_teacher_name(subject_key, teacher_key)}\n\n"
        f"{current_text}\n\n"
        f"👍 {like_label(user_id)}: {like_count} ta ({like_percent:.1f}%)\n"
        f"👎 {dislike_label(user_id)}: {dislike_count} ta ({dislike_percent:.1f}%)\n"
        f"Jami: {total} ta\n\n"
        f"Bahoni tanlang yoki o'zgartiring:"
    )

def get_complaint_intro_text(user_id: int) -> str:
    return tr(
        user_id,
        "📩 <b>Shikoyat va takliflar</b>\n\n"
        "Shikoyat, taklif yoki murojaatingizni bitta xabar qilib yozing.\n"
        "Foydalanuvchi ma'lumotlari sir tutiladi.\n\n"
        "Bekor qilish uchun <b>❌ Bekor qilish</b> tugmasini bosing."
    )


def get_complaint_saved_text(user_id: int) -> str:
    return tr(user_id, "✅ <b>Xabaringiz qabul qilindi.</b>\n\nRahmat, murojaatingiz adminlarga yuborildi.")

def get_results_menu_text(user_id: int, is_admin_view: bool = False) -> str:
    title = "Admin natijalar bo'limi" if is_admin_view else "Natijalar bo'limi"
    return tr(user_id, f"📊 <b>{title}</b>\n\nKerakli bo'limni tanlang:")

def get_admin_panel_text(user_id: int) -> str:
    status_text = "🟢 Ochiq" if is_voting_open() else "🔴 Yopiq"
    return tr(user_id, f"🎛 <b>Admin panel</b>\n\nVoting holati: {status_text}\nJami ovozlar: {get_total_votes()}")

def get_general_results_text(user_id: int) -> str:
    cursor.execute("""
        SELECT subject_key, teacher_key, COUNT(*)
        FROM votes
        GROUP BY subject_key, teacher_key
    """)
    counts = {}
    for subject_key, teacher_key, count in cursor.fetchall():
        counts[(normalize_subject_key(subject_key), teacher_key)] = count

    total_votes = sum(counts.values())
    lines = ["📊 <b>Umumiy natijalar</b>\n"]

    for subject_key, teacher_key, teacher_name in get_all_teachers_flat():
        count = counts.get((subject_key, teacher_key), 0)
        percent = get_vote_percent(count, total_votes)
        lines.append(
            f"<b>{teacher_name}</b> — {get_subject_name(subject_key)}\n"
            f"<code>{build_progress_bar(percent)}</code>  <b>{percent:.1f}%</b>  •  {count} ta\n"
        )

    lines.append(f"🗳 <b>Jami ovozlar:</b> {total_votes}")
    lines.append(f"{'🟢' if is_voting_open() else '🔴'} <b>Holat:</b> {'Ochiq' if is_voting_open() else 'Yopiq'}")

    text = "\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n\n... qisqartirildi"
    return tr(user_id, text)


def get_subject_results_text(user_id: int, subject_key: str) -> str:
    subject_key = normalize_subject_key(subject_key)
    if subject_key not in SUBJECTS:
        return tr(user_id, "Noto'g'ri kafedra.")

    cursor.execute("""
        SELECT teacher_key, COUNT(*)
        FROM votes
        WHERE subject_key = ?
        GROUP BY teacher_key
    """, (subject_key,))
    subject_counts = {teacher_key: count for teacher_key, count in cursor.fetchall()}

    cursor.execute("SELECT COUNT(*) FROM votes")
    total_votes = cursor.fetchone()[0]
    subject_total = sum(subject_counts.values())

    lines = [f"📊 <b>{get_subject_name(subject_key)} bo'yicha natijalar</b>\n"]

    for teacher_key, teacher_name in SUBJECTS[subject_key]["teachers"].items():
        count = subject_counts.get(teacher_key, 0)
        percent = get_vote_percent(count, subject_total)
        lines.append(
            f"<b>{teacher_name}</b>\n"
            f"<code>{build_progress_bar(percent)}</code>  <b>{percent:.1f}%</b>  •  {count} ta\n"
        )

    lines.append(f"🗳 <b>Ushbu kafedra ovozlari:</b> {subject_total}")
    lines.append(f"🗳 <b>Jami ovozlar:</b> {total_votes}")
    lines.append(f"{'🟢' if is_voting_open() else '🔴'} <b>Holat:</b> {'Ochiq' if is_voting_open() else 'Yopiq'}")

    text = "\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n\n... qisqartirildi"
    return tr(user_id, text)


# FIX #9: Optional ishlatildi
def get_rating_stats_text(user_id: int, subject_key: Optional[str] = None) -> str:
    rows = rating_rows()
    if subject_key and subject_key != "general":
        subject_key = normalize_subject_key(subject_key)
        rows = [r for r in rows if r["subject_key"] == subject_key]
        title = f"⭐️ <b>{get_subject_name(subject_key)} — baholash foizlari</b>\n"
    else:
        title = "⭐️ <b>Umumiy baholash foizlari</b>\n"
    lines = [title]
    for r in rows:
        lines.append(
            f"<b>{r['teacher_name']}</b> — {r['subject_name']}\n"
            f"👍 {r['like']} ta ({r['like_percent']:.1f}%) | "
            f"👎 {r['dislike']} ta ({r['dislike_percent']:.1f}%) | "
            f"Jami: {r['total']}\n"
        )
    text = "\n".join(lines)
    return tr(user_id, text[:4000] + ("\n\n... qisqartirildi" if len(text) > 4000 else ""))

def get_top_ratings_text(user_id: int) -> str:
    rows = [r for r in rating_rows() if r["total"] > 0]

    def line_items(items, percent_key: str, icon: str):
        if not items:
            return "Ma'lumot yo'q"
        return "\n".join([
            f"{i}. {r['teacher_name']} — {r['subject_name']} | {icon} {r[percent_key]:.1f}% | Jami: {r['total']}"
            for i, r in enumerate(items, 1)
        ])

    high_like = sorted(rows, key=lambda r: (r["like_percent"], r["total"]), reverse=True)[:10]
    high_dislike = sorted(rows, key=lambda r: (r["dislike_percent"], r["total"]), reverse=True)[:10]

    text = (
        "🏆 <b>TOP reytinglar</b>\n\n"
        "🔝 <b>TOP 10 eng baland like nisbati</b>\n" + line_items(high_like, "like_percent", "👍") + "\n\n"
        "🔻 <b>TOP 10 eng baland dislike nisbati</b>\n" + line_items(high_dislike, "dislike_percent", "👎")
    )
    return tr(user_id, text[:4000] + ("\n\n... qisqartirildi" if len(text) > 4000 else ""))


def get_users_text(user_id: int) -> str:
    cursor.execute("SELECT user_id, full_name, username, subject_key, teacher_key, voted_at FROM votes ORDER BY voted_at DESC")
    rows = cursor.fetchall()
    if not rows:
        return tr(user_id, "👥 Hali hech kim ovoz bermagan.")
    lines = [f"👥 <b>Kim kimga ovoz berdi</b>\n\nJami: {len(rows)} ta foydalanuvchi\n"]
    for i, (uid, full_name, username, subject_key, teacher_key, voted_at) in enumerate(rows, start=1):
        subject_key = normalize_subject_key(subject_key)
        # FIX #3: f-string ichida apostrof xatosi — o'zgaruvchiga olindi
        name = full_name or "Noma'lum"
        line = f"{i}. <b>{name}</b>"
        if username:
            line += f" (@{username})"
        line += f"\n   → Kafedra: {get_subject_name(subject_key)}"
        line += f"\n   → O'qituvchi: {get_teacher_name(subject_key, teacher_key)}"
        line += f"\n   → ID: <code>{uid}</code>"
        if voted_at:
            line += f"\n   → {voted_at}"
        lines.append(line)
    text = "\n\n".join(lines)
    return tr(user_id, text[:4000] + ("\n\n... qisqartirildi" if len(text) > 4000 else ""))


def get_my_vote_text(user_id: int) -> str:
    cursor.execute("""
        SELECT subject_key, teacher_key, voted_at
        FROM votes
        WHERE user_id = ?
    """, (user_id,))
    row = cursor.fetchone()
    if not row:
        return tr(user_id, "🧾 <b>Mening ovozim</b>\n\nSiz hali asosiy ovoz bermagansiz.")

    subject_key, teacher_key, voted_at = row
    subject_key = normalize_subject_key(subject_key)
    return tr(
        user_id,
        f"🧾 <b>Mening ovozim</b>\n\n"
        f"<b>Kafedra:</b> {get_subject_name(subject_key)}\n"
        f"<b>O'qituvchi:</b> {get_teacher_name(subject_key, teacher_key)}\n"
        f"<b>Sana:</b> {voted_at or 'Noma\'lum'}"
    )


def get_my_ratings_text(user_id: int) -> str:
    cursor.execute("""
        SELECT subject_key, teacher_key, rating, rated_at
        FROM teacher_ratings
        WHERE user_id = ?
        ORDER BY rated_at DESC
    """, (user_id,))
    rows = cursor.fetchall()
    if not rows:
        return tr(user_id, "⭐️ <b>Mening baholarim</b>\n\nSiz hali o'qituvchilarga like/dislike bermagansiz.")

    lines = [f"⭐️ <b>Mening baholarim</b>\n\nJami: {len(rows)} ta baho\n"]
    for i, (subject_key, teacher_key, rating, rated_at) in enumerate(rows, start=1):
        subject_key = normalize_subject_key(subject_key)
        icon = "👍" if rating == "like" else "👎"
        label = like_label(user_id) if rating == "like" else dislike_label(user_id)
        lines.append(
            f"{i}. <b>{get_teacher_name(subject_key, teacher_key)}</b>\n"
            f"   Kafedra: {get_subject_name(subject_key)}\n"
            f"   Baho: {icon} {label}\n"
            f"   Sana: {rated_at or ''}"
        )

    text = "\n\n".join(lines)
    return tr(user_id, text[:4000] + ("\n\n... qisqartirildi" if len(text) > 4000 else ""))


def get_teacher_detailed_stats_text(user_id: int, subject_key: str, teacher_key: str) -> str:
    subject_key = normalize_subject_key(subject_key)
    if subject_key not in SUBJECTS or teacher_key not in SUBJECTS[subject_key]["teachers"]:
        return tr(user_id, "Noto'g'ri o'qituvchi tanlandi.")

    cursor.execute("SELECT COUNT(*) FROM votes WHERE subject_key = ? AND teacher_key = ?", (subject_key, teacher_key))
    vote_count = cursor.fetchone()[0]
    subject_total = get_total_votes(subject_key)
    total_votes = get_total_votes()
    vote_percent_subject = get_vote_percent(vote_count, subject_total)
    vote_percent_total = get_vote_percent(vote_count, total_votes)
    like_count, dislike_count, rating_total, like_percent, dislike_percent = get_rating_counts(subject_key, teacher_key)

    return tr(
        user_id,
        f"👤 <b>O'qituvchi statistikasi</b>\n\n"
        f"<b>O'qituvchi:</b> {get_teacher_name(subject_key, teacher_key)}\n"
        f"<b>Kafedra:</b> {get_subject_name(subject_key)}\n\n"
        f"🗳 <b>Asosiy ovozlar:</b> {vote_count} ta\n"
        f"📊 <b>Kafedra ichidagi ulushi:</b> {vote_percent_subject:.1f}%\n"
        f"🌐 <b>Jami ovozlar ichidagi ulushi:</b> {vote_percent_total:.1f}%\n\n"
        f"👍 <b>{like_label(user_id)}:</b> {like_count} ta ({like_percent:.1f}%)\n"
        f"👎 <b>{dislike_label(user_id)}:</b> {dislike_count} ta ({dislike_percent:.1f}%)\n"
        f"⭐️ <b>Jami baholar:</b> {rating_total} ta"
    )


def create_backup_zip() -> str:
    conn.commit()
    votes_path = export_votes_to_excel()
    rating_path = export_rating_to_excel()
    complaints_path = export_complaints_to_docx()

    with zipfile.ZipFile(BACKUP_ZIP_FILE, "w", zipfile.ZIP_DEFLATED) as zf:
        if os.path.exists(DB_NAME):
            zf.write(DB_NAME, arcname="votes.db")
        for path in (votes_path, rating_path, complaints_path):
            if path and os.path.exists(path):
                zf.write(path, arcname=os.path.basename(path))
    return BACKUP_ZIP_FILE


def get_results_text_by_scope(user_id: int, scope: str) -> str:
    scope = normalize_subject_key(scope)
    if scope == "general":
        return get_general_results_text(user_id)
    if scope in SUBJECTS:
        return get_subject_results_text(user_id, scope)
    return tr(user_id, "Noto'g'ri bo'lim.")

# =========================
# EXPORT
# =========================
def export_votes_to_csv() -> str:
    cursor.execute("SELECT user_id, full_name, username, subject_key, teacher_key, voted_at FROM votes ORDER BY voted_at DESC")
    rows = cursor.fetchall()
    with open(EXPORT_FILE, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["User ID", "Full Name", "Username", "Subject", "Teacher", "Voted At"])
        for user_id, full_name, username, subject_key, teacher_key, voted_at in rows:
            subject_key = normalize_subject_key(subject_key)
            writer.writerow([user_id, full_name or "", username or "", get_subject_name(subject_key), get_teacher_name(subject_key, teacher_key), voted_at or ""])
    return EXPORT_FILE

def ws_append_header(ws, headers):
    ws.append(headers)
    for cell in ws[1]:
        cell.style = "Headline 4"

def export_votes_to_excel() -> str:
    if Workbook is None:
        return export_votes_to_csv()
    wb = Workbook()
    wb.remove(wb.active)

    ws = wb.create_sheet("Umumiy ovozlar")
    ws_append_header(ws, ["User ID", "Full Name", "Username", "Kafedra", "O'qituvchi", "Voted At"])
    cursor.execute("SELECT user_id, full_name, username, subject_key, teacher_key, voted_at FROM votes ORDER BY voted_at DESC")
    for user_id, full_name, username, subject_key, teacher_key, voted_at in cursor.fetchall():
        subject_key = normalize_subject_key(subject_key)
        ws.append([user_id, full_name or "", username or "", get_subject_name(subject_key), get_teacher_name(subject_key, teacher_key), voted_at or ""])

    for subject_key, subject_data in SUBJECTS.items():
        ws = wb.create_sheet(subject_data["name"][:31])
        ws_append_header(ws, ["User ID", "Full Name", "Username", "O'qituvchi", "Voted At"])
        cursor.execute("SELECT user_id, full_name, username, teacher_key, voted_at FROM votes WHERE subject_key = ? ORDER BY voted_at DESC", (subject_key,))
        for user_id, full_name, username, teacher_key, voted_at in cursor.fetchall():
            ws.append([user_id, full_name or "", username or "", get_teacher_name(subject_key, teacher_key), voted_at or ""])

    ws = wb.create_sheet("Umumiy natija")
    ws_append_header(ws, ["Kafedra", "O'qituvchi", "Ovozlar", "Foiz"])
    total = get_total_votes()
    for subject_key, teacher_key, teacher_name in get_all_teachers_flat():
        cursor.execute("SELECT COUNT(*) FROM votes WHERE subject_key = ? AND teacher_key = ?", (subject_key, teacher_key))
        count = cursor.fetchone()[0]
        ws.append([get_subject_name(subject_key), teacher_name, count, round((count / total * 100) if total else 0, 2)])

    for subject_key, subject_data in SUBJECTS.items():
        ws = wb.create_sheet((subject_data["name"][:24] + " natija")[:31])
        ws_append_header(ws, ["O'qituvchi", "Ovozlar", "Kafedra ichidagi foiz"])
        subject_total = get_total_votes(subject_key)
        for teacher_key, teacher_name in subject_data["teachers"].items():
            cursor.execute("SELECT COUNT(*) FROM votes WHERE subject_key = ? AND teacher_key = ?", (subject_key, teacher_key))
            count = cursor.fetchone()[0]
            ws.append([teacher_name, count, round((count / subject_total * 100) if subject_total else 0, 2)])

    wb.save(VOTES_XLSX_FILE)
    return VOTES_XLSX_FILE

def export_rating_to_excel() -> str:
    if Workbook is None:
        path = os.path.join(DATA_DIR, "rating_export.csv")
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["Kafedra", "O'qituvchi", "Like", "Dislike", "Jami", "Like %", "Dislike %"])
            for r in rating_rows():
                writer.writerow([r["subject_name"], r["teacher_name"], r["like"], r["dislike"], r["total"], round(r["like_percent"], 2), round(r["dislike_percent"], 2)])
        return path

    wb = Workbook()
    wb.remove(wb.active)

    ws = wb.create_sheet("Umumiy rating")
    ws_append_header(ws, ["Kafedra", "O'qituvchi", "Like", "Dislike", "Jami", "Like %", "Dislike %"])
    for r in rating_rows():
        ws.append([r["subject_name"], r["teacher_name"], r["like"], r["dislike"], r["total"], round(r["like_percent"], 2), round(r["dislike_percent"], 2)])

    for subject_key, subject_data in SUBJECTS.items():
        ws = wb.create_sheet((subject_data["name"][:24] + " rating")[:31])
        ws_append_header(ws, ["O'qituvchi", "Like", "Dislike", "Jami", "Like %", "Dislike %"])
        for r in [x for x in rating_rows() if x["subject_key"] == subject_key]:
            ws.append([r["teacher_name"], r["like"], r["dislike"], r["total"], round(r["like_percent"], 2), round(r["dislike_percent"], 2)])

    ws = wb.create_sheet("Umumiy ovozlar")
    ws_append_header(ws, ["User ID", "Full Name", "Username", "Kafedra", "O'qituvchi", "Rating", "Rated At"])
    cursor.execute("SELECT user_id, full_name, username, subject_key, teacher_key, rating, rated_at FROM teacher_ratings ORDER BY rated_at DESC")
    for user_id, full_name, username, subject_key, teacher_key, rating, rated_at in cursor.fetchall():
        subject_key = normalize_subject_key(subject_key)
        ws.append([user_id, full_name or "", username or "", get_subject_name(subject_key), get_teacher_name(subject_key, teacher_key), rating, rated_at or ""])

    wb.save(RATING_XLSX_FILE)
    return RATING_XLSX_FILE


def export_complaints_to_word() -> str:
    cursor.execute("""
        SELECT user_id, full_name, username, message_text, created_at
        FROM complaints
        ORDER BY created_at DESC
    """)
    rows = cursor.fetchall()

    if Document is None:
        txt_path = os.path.join(DATA_DIR, "shikoyat_takliflar.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("Shikoyat va takliflar\n")
            f.write("=" * 30 + "\n\n")
            if not rows:
                f.write("Hali shikoyat yoki taklif yo'q.\n")
            for i, (user_id, full_name, username, message_text, created_at) in enumerate(rows, 1):
                # FIX #3: f-string ichida apostrof xatosi tuzatildi
                name = full_name or "Noma'lum"
                uname = f"@{username}" if username else "yo'q"
                f.write(f"{i}. Foydalanuvchi: {name}\n")
                f.write(f"   Username: {uname}\n")
                f.write(f"   ID: {user_id}\n")
                f.write(f"   Sana: {created_at or ''}\n")
                f.write(f"   Matn: {message_text or ''}\n")
                f.write("-" * 30 + "\n")
        return txt_path

    doc = Document()
    doc.add_heading("Shikoyat va takliflar", level=1)

    if not rows:
        doc.add_paragraph("Hali shikoyat yoki taklif yo'q.")
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["№", "F.I.Sh", "Username", "User ID", "Sana", "Matn"]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        for i, (user_id, full_name, username, message_text, created_at) in enumerate(rows, 1):
            cells = table.add_row().cells
            cells[0].text = str(i)
            cells[1].text = full_name or "Noma'lum"
            cells[2].text = f"@{username}" if username else ""
            cells[3].text = str(user_id)
            cells[4].text = created_at or ""
            cells[5].text = message_text or ""

    doc.save(COMPLAINTS_DOCX_FILE)
    return COMPLAINTS_DOCX_FILE

# =========================
# SUBSCRIPTION
# =========================
async def check_user_subscription(user_id: int) -> bool:
    try:
        member = await bot.get_chat_member(CHANNEL_USERNAME, user_id)
        if member.status in {ChatMemberStatus.CREATOR, ChatMemberStatus.ADMINISTRATOR, ChatMemberStatus.MEMBER}:
            return True
        if member.status == ChatMemberStatus.RESTRICTED:
            return bool(getattr(member, "is_member", False))
        return False
    except Exception as e:
        logging.error(f"Obunani tekshirishda xatolik: {e}")
        return False

async def safe_edit_message(callback: CallbackQuery, text: str, reply_markup: Optional[InlineKeyboardMarkup] = None):
    try:
        await callback.message.edit_text(text=text, parse_mode="HTML", reply_markup=reply_markup)
    except TelegramBadRequest as e:
        if "message is not modified" in str(e).lower():
            return
        logging.error(f"edit_text xatosi: {e}")
        try:
            await callback.message.answer(text=text, parse_mode="HTML", reply_markup=reply_markup)
        except Exception as e2:
            logging.error(f"answer fallback xatosi: {e2}")
    except Exception as e:
        logging.error(f"safe_edit_message umumiy xato: {e}")



def add_refresh_time(text: str, user_id: int) -> str:
    """Telegram edit_text 'message is not modified' xatosini oldini olish uchun vaqt qo'shadi."""
    return text + tr(user_id, f"  ⏱ Yangilandi: {uz_now().strftime('%H:%M:%S.%f')[:-3]}")


def can_start_refresh(user_id: int, key: str) -> bool:
    """
    Bitta user bitta bo'limda refreshni ketma-ket bosib yuborsa,
    parallel callbacklar natijani chalkashtirib yubormasligi uchun cheklaydi.
    """
    now = uz_now().timestamp()
    refresh_key = (user_id, key)

    if refresh_key in REFRESH_BUSY:
        return False

    last = LAST_REFRESH.get(refresh_key, 0)
    if now - last < REFRESH_COOLDOWN_SECONDS:
        return False

    REFRESH_BUSY.add(refresh_key)
    LAST_REFRESH[refresh_key] = now
    return True


def finish_refresh(user_id: int, key: str):
    REFRESH_BUSY.discard((user_id, key))


def get_settings_text(user_id: int) -> str:
    script = get_user_script(user_id)
    current = "Lotin" if script == "latin" else "Крилл"
    return tr(
        user_id,
        f"⚙️ <b>Sozlamalar</b>\n\n"
        f"Hozirgi yozuv: <b>{current}</b>\n\n"
        f"Kerakli yozuv turini tanlang:"
    )

# =========================
# KEYBOARDS
# =========================
def subscription_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text="📸 Instagram", url=INSTAGRAM_URL))
    kb.row(InlineKeyboardButton(text="📘 Facebook", url=FACEBOOK_URL))
    kb.row(InlineKeyboardButton(text="📢 Telegram", url=CHANNEL_URL))
    kb.row(InlineKeyboardButton(text="✅ Tekshirish", callback_data="check_subscription"))
    return kb.as_markup()

def home_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    if has_access(user_id):
        kb.row(
            InlineKeyboardButton(text=tr(user_id, "🗳 Ovoz berish"), callback_data="go_vote_panel"),
            InlineKeyboardButton(text=tr(user_id, "⭐️ O'qituvchilarni baholash"), callback_data="go_rating_panel")
        )
        kb.row(InlineKeyboardButton(text=tr(user_id, "📊 Natijalar"), callback_data="show_results_menu_user"))
        kb.row(
            InlineKeyboardButton(text=tr(user_id, "🧾 Mening ovozim"), callback_data="my_vote"),
            InlineKeyboardButton(text=tr(user_id, "⭐️ Mening baholarim"), callback_data="my_ratings")
        )
        kb.row(InlineKeyboardButton(text=tr(user_id, "📩 Shikoyat va takliflar"), callback_data="go_complaint_panel"))
    else:
        kb.row(InlineKeyboardButton(text=tr(user_id, "✅ Obunani tekshirish"), callback_data="check_subscription"))

    kb.row(
        InlineKeyboardButton(text=tr(user_id, "ℹ️ Yordam"), callback_data="help_info"),
        InlineKeyboardButton(text=tr(user_id, "⚙️ Sozlamalar"), callback_data="user_settings")
    )
    return kb.as_markup()

def settings_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    current = get_user_script(user_id)
    latin_text = "✅ Lotin" if current == "latin" else "Lotin"
    cyrillic_text = "✅ Крилл" if current == "cyrillic" else "Крилл"
    kb.row(
        InlineKeyboardButton(text=latin_text, callback_data="set_script:latin"),
        InlineKeyboardButton(text=cyrillic_text, callback_data="set_script:cyrillic")
    )
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_home"))
    return kb.as_markup()

def subjects_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in SUBJECTS.items():
        kb.row(InlineKeyboardButton(text=tr(user_id, subject_data["name"]), callback_data=f"subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()

def teachers_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    teachers = list(SUBJECTS[subject_key]["teachers"].items())
    for i in range(0, len(teachers), 2):
        row = [InlineKeyboardButton(text=tr(user_id, teacher_name), callback_data=f"vote:{subject_key}:{teacher_key}") for teacher_key, teacher_name in teachers[i:i + 2]]
        kb.row(*row)
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Kafedralarga qaytish"), callback_data="go_vote_panel"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()

def rating_subjects_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in SUBJECTS.items():
        kb.row(InlineKeyboardButton(text=tr(user_id, subject_data["name"]), callback_data=f"rating_subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()

def rating_teachers_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    teachers = list(SUBJECTS[subject_key]["teachers"].items())
    for i in range(0, len(teachers), 2):
        row = [InlineKeyboardButton(text=tr(user_id, teacher_name), callback_data=f"rating_teacher:{subject_key}:{teacher_key}") for teacher_key, teacher_name in teachers[i:i + 2]]
        kb.row(*row)
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Kafedralarga qaytish"), callback_data="go_rating_panel"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()

def rate_keyboard(user_id: int, subject_key: str, teacher_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=f"👍 {like_label(user_id)}", callback_data=f"rate:like:{subject_key}:{teacher_key}"),
        InlineKeyboardButton(text=f"👎 {dislike_label(user_id)}", callback_data=f"rate:dislike:{subject_key}:{teacher_key}")
    )
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ O'qituvchilar"), callback_data=f"rating_subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()

def results_menu_keyboard_user(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in SUBJECTS.items():
        kb.row(InlineKeyboardButton(text=tr(user_id, subject_data["name"]), callback_data=f"show_results_user:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_home"))
    return kb.as_markup()

def results_menu_keyboard_admin(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in SUBJECTS.items():
        kb.row(InlineKeyboardButton(text=subject_data["name"], callback_data=f"show_results_admin:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()

def rating_results_menu_keyboard_admin(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in SUBJECTS.items():
        kb.row(InlineKeyboardButton(text=subject_data["name"], callback_data=f"show_rating_stats:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()

def results_keyboard_user(user_id: int, scope: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data=f"refresh_results_user:{scope}"),
        InlineKeyboardButton(text=tr(user_id, "📂 Bo'limlar"), callback_data="show_results_menu_user")
    )
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()

def results_keyboard_admin(user_id: int, scope: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data=f"refresh_results_admin:{scope}"),
        InlineKeyboardButton(text=tr(user_id, "📂 Bo'limlar"), callback_data="admin_results")
    )
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()

def rating_stats_keyboard_admin(user_id: int, scope: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data=f"refresh_rating_stats:{scope}"),
        InlineKeyboardButton(text=tr(user_id, "📂 Bo'limlar"), callback_data="admin_rating_stats")
    )
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()

def after_vote_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "📊 Natijalar"), callback_data="show_results_menu_user"),
        InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home")
    )
    return kb.as_markup()


def confirm_vote_keyboard(user_id: int, subject_key: str, teacher_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "✅ Tasdiqlash"), callback_data=f"confirm_vote:{subject_key}:{teacher_key}"),
        InlineKeyboardButton(text=tr(user_id, "❌ Bekor qilish"), callback_data=f"subject:{subject_key}")
    )
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ O'qituvchilar"), callback_data=f"subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def simple_back_home_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()


def teacher_stats_subjects_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    for subject_key, subject_data in SUBJECTS.items():
        kb.row(InlineKeyboardButton(text=tr(user_id, subject_data["name"]), callback_data=f"teacher_stats_subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def teacher_stats_teachers_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    teachers = list(SUBJECTS[subject_key]["teachers"].items())
    for i in range(0, len(teachers), 2):
        row = [InlineKeyboardButton(text=tr(user_id, teacher_name), callback_data=f"teacher_stats:{subject_key}:{teacher_key}") for teacher_key, teacher_name in teachers[i:i + 2]]
        kb.row(*row)
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Kafedralarga qaytish"), callback_data="admin_teacher_stats"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()


def teacher_stats_keyboard(user_id: int, subject_key: str) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ O'qituvchilar"), callback_data=f"teacher_stats_subject:{subject_key}"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()

def admin_panel_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text="📊 Ovoz natijalari", callback_data="admin_results"))
    kb.row(InlineKeyboardButton(text="⭐️ Baholash foizlari", callback_data="admin_rating_stats"))
    kb.row(InlineKeyboardButton(text="🏆 TOP reytinglar", callback_data="admin_top_ratings"))
    kb.row(InlineKeyboardButton(text="👥 Foydalanuvchilar", callback_data="admin_users"))
    kb.row(InlineKeyboardButton(text="👤 O'qituvchi statistikasi", callback_data="admin_teacher_stats"))
    kb.row(InlineKeyboardButton(text="📩 Shikoyat/takliflar", callback_data="admin_complaints"))
    kb.row(
        InlineKeyboardButton(text="📁 Excel ovozlar", callback_data="admin_export_votes_excel"),
        InlineKeyboardButton(text="📁 Excel rating", callback_data="admin_export_rating_excel")
    )
    kb.row(
        InlineKeyboardButton(text="🔓 Open", callback_data="admin_open"),
        InlineKeyboardButton(text="🔒 Close", callback_data="admin_close")
    )
    kb.row(
        InlineKeyboardButton(text="♻️ Reset ovozlar", callback_data="admin_reset_votes_confirm"),
        InlineKeyboardButton(text="🗑 Reset rating", callback_data="admin_reset_rating_confirm")
    )
    kb.row(InlineKeyboardButton(text="🧹 Shikoyatlarni tozalash", callback_data="admin_reset_complaints_confirm"))
    kb.row(InlineKeyboardButton(text="💾 Backup", callback_data="admin_backup"))
    return kb.as_markup()

def reset_confirm_keyboard(user_id: int, mode: str = "votes") -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    if mode == "votes":
        yes_cb = "admin_reset_votes"
    elif mode == "rating":
        yes_cb = "admin_reset_rating"
    elif mode == "complaints":
        yes_cb = "admin_reset_complaints"
    else:
        yes_cb = "admin_reset_votes"
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "❌ Bekor qilish"), callback_data="cancel_reset"),
        InlineKeyboardButton(text="✅ Ha, o'chirish", callback_data=yes_cb)
    )
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()

def complaint_cancel_keyboard(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "❌ Bekor qilish"), callback_data="cancel_complaint"))
    kb.row(InlineKeyboardButton(text=tr(user_id, "🏠 Bosh menyu"), callback_data="go_home"))
    return kb.as_markup()

def complaints_keyboard_admin(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data="refresh_admin_complaints"),
        InlineKeyboardButton(text="📄 Word", callback_data="admin_export_complaints_docx")
    )
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    return kb.as_markup()

def users_keyboard_admin(user_id: int) -> InlineKeyboardMarkup:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data="refresh_admin_users"),
        InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel")
    )
    return kb.as_markup()

# =========================
# START / COMMANDS
# =========================
@dp.message(Command("start"))
async def start_handler(message: Message):
    user_id = message.from_user.id
    ensure_user(user_id)

    if has_access(user_id):
        await message.answer(get_home_text(user_id), parse_mode="HTML", reply_markup=home_keyboard(user_id))
        return

    if not await check_user_subscription(user_id):
        await message.answer(get_welcome_text(user_id), parse_mode="HTML", reply_markup=subscription_keyboard(user_id))
        return

    grant_access(user_id)
    await message.answer(get_home_text(user_id), parse_mode="HTML", reply_markup=home_keyboard(user_id))


@dp.message(Command("my_access"))
async def my_access_handler(message: Message):
    user_id = message.from_user.id
    ensure_user(user_id)
    if not is_admin(user_id):
        return
    cursor.execute("SELECT access_granted FROM user_prefs WHERE user_id = ?", (user_id,))
    row = cursor.fetchone()
    # FIX #3: apostrof xatosi tuzatildi
    val = row[0] if row else "yo'q"
    await message.answer(f"access_granted: {val}")

@dp.message(Command("check_channel"))
async def check_channel_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    try:
        me = await bot.get_me()
        member = await bot.get_chat_member(CHANNEL_USERNAME, me.id)
        if member.status in {ChatMemberStatus.ADMINISTRATOR, ChatMemberStatus.CREATOR}:
            await message.answer(f"✅ Bot kanalni ko'ra olyapti va admin.\nKanal: {CHANNEL_USERNAME}")
        else:
            await message.answer("⚠️ Bot kanalni ko'ryapti, lekin admin emas. Obuna tekshiruvi to'liq ishlashi uchun botni kanalga admin qiling.")
    except Exception as e:
        await message.answer(f"❌ Kanalni tekshirib bo'lmadi.\nBotni kanalga admin qiling.\nXato: {e}")

@dp.message(Command("results"))
async def results_handler(message: Message):
    user_id = message.from_user.id
    ensure_user(user_id)
    await message.answer(get_results_menu_text(user_id, False), parse_mode="HTML", reply_markup=results_menu_keyboard_user(user_id))

@dp.message(Command("debug_eshnazarova"))
async def debug_eshnazarova_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        return

    cursor.execute("""
        SELECT user_id, full_name, username, subject_key, teacher_key, voted_at
        FROM votes
        WHERE teacher_key = 'aif_10'
        ORDER BY voted_at DESC
    """)
    rows = cursor.fetchall()
    if not rows:
        await message.answer("Eshnazarova Maziya Allanazarovna uchun bazada ovoz yo'q.")
        return

    lines = ["Eshnazarova Maziya Allanazarovna uchun bazadagi ovozlar:"]
    for uid, full_name, username, subject_key, teacher_key, voted_at in rows:
        lines.append(f"ID: {uid} | {full_name or ''} | @{username or ''} | {subject_key}/{teacher_key} | {voted_at}")
    await message.answer("\n".join(lines[:50]))

@dp.message(Command("admin"))
async def admin_panel_handler(message: Message):
    user_id = message.from_user.id
    ensure_user(user_id)
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    await message.answer(get_admin_panel_text(user_id), parse_mode="HTML", reply_markup=admin_panel_keyboard(user_id))

@dp.message(Command("users"))
async def admin_users_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    await message.answer(get_users_text(user_id), parse_mode="HTML", reply_markup=users_keyboard_admin(user_id))

@dp.message(Command("export"))
async def admin_export_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    filename = export_votes_to_excel()
    await message.answer_document(FSInputFile(filename), caption="📁 Ovozlar Excel fayl ko'rinishida.")

@dp.message(Command("open"))
async def admin_open_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    open_voting()
    await message.answer("🟢 Ovoz berish ochildi.")

@dp.message(Command("close"))
async def admin_close_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    close_voting()
    await message.answer("🔴 Ovoz berish yopildi.")

@dp.message(Command("reset_votes"))
async def admin_reset_handler(message: Message):
    user_id = message.from_user.id
    if not is_admin(user_id):
        await message.answer("Siz admin emassiz.")
        return
    await message.answer(
        "⚠️ <b>Diqqat!</b>\n\nBarcha ovozlar o'chiriladi.\nDavom etasizmi?",
        parse_mode="HTML",
        reply_markup=reset_confirm_keyboard(user_id, "votes")
    )

# =========================
# USER CALLBACKS
# =========================
@dp.callback_query(F.data == "go_home")
async def go_home_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    ensure_user(user_id)
    if require_access_only(user_id):
        await safe_edit_message(callback, get_home_text(user_id), home_keyboard(user_id))
    else:
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "help_info")
async def help_info_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "⬅️ Orqaga"), callback_data="go_home"))
    await safe_edit_message(callback, get_help_text(user_id), kb.as_markup())
    await callback.answer()

@dp.callback_query(F.data == "my_vote")
async def my_vote_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    await safe_edit_message(callback, get_my_vote_text(user_id), simple_back_home_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "my_ratings")
async def my_ratings_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    await safe_edit_message(callback, get_my_ratings_text(user_id), simple_back_home_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "go_complaint_panel")
async def go_complaint_panel_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    ensure_user(user_id)

    if not has_access(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return

    WAITING_COMPLAINT_TEXT.add(user_id)
    await safe_edit_message(callback, get_complaint_intro_text(user_id), complaint_cancel_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data == "cancel_complaint")
async def cancel_complaint_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    WAITING_COMPLAINT_TEXT.discard(user_id)

    if has_access(user_id):
        await safe_edit_message(callback, get_home_text(user_id), home_keyboard(user_id))
    else:
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
    await callback.answer(tr(user_id, "Bekor qilindi"))


@dp.callback_query(F.data == "user_settings")
async def user_settings_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    ensure_user(user_id)
    await safe_edit_message(callback, get_settings_text(user_id), settings_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data.startswith("set_script:"))
async def set_script_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    ensure_user(user_id)

    script = callback.data.split(":", 1)[1]
    if script not in ("latin", "cyrillic"):
        await callback.answer("Xato", show_alert=True)
        return

    set_user_script(user_id, script)
    await safe_edit_message(callback, get_settings_text(user_id), settings_keyboard(user_id))
    await callback.answer(tr(user_id, "Yozuv turi saqlandi"))

@dp.callback_query(F.data == "check_subscription")
async def check_subscription_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    ensure_user(user_id)

    ok = await check_user_subscription(user_id)
    if not ok:
        reset_access(user_id)
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return

    grant_access(user_id)

    # FIX #5: grant_access dan keyin keraksiz tekshiruv olib tashlandi
    await safe_edit_message(
        callback,
        "✅ <b>Obuna tasdiqlandi</b>\n\nEndi bosh menyudan bemalol foydalanishingiz mumkin:",
        home_keyboard(user_id)
    )
    await callback.answer(tr(user_id, "Tasdiqlandi"))


@dp.callback_query(F.data == "go_vote_panel")
async def go_vote_panel_handler(callback: CallbackQuery):
    user_id = callback.from_user.id

    # Ovoz berishdan oldin obuna holati qayta tekshiriladi.
    if not await check_user_subscription(user_id):
        reset_access(user_id)
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return

    grant_access(user_id)
    if has_voted(user_id):
        await safe_edit_message(callback, get_already_voted_text(user_id), home_keyboard(user_id))
        await callback.answer()
        return
    if not is_voting_open():
        await safe_edit_message(callback, get_closed_text(user_id), home_keyboard(user_id))
        await callback.answer()
        return
    await safe_edit_message(callback, get_subject_select_text(user_id), subjects_keyboard(user_id))
    await callback.answer()

@dp.callback_query(F.data.startswith("subject:"))
async def subject_select_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    if has_voted(user_id):
        await safe_edit_message(callback, get_already_voted_text(user_id), home_keyboard(user_id))
        await callback.answer()
        return
    if not is_voting_open():
        await safe_edit_message(callback, get_closed_text(user_id), home_keyboard(user_id))
        await callback.answer()
        return
    subject_key = normalize_subject_key(callback.data.split(":")[1])
    if subject_key not in SUBJECTS:
        await callback.answer("Noto'g'ri bo'lim tanlandi.", show_alert=True)
        return
    await safe_edit_message(callback, get_teacher_select_text(user_id, subject_key), teachers_keyboard(user_id, subject_key))
    await callback.answer()

@dp.callback_query(F.data.startswith("vote:"))
async def vote_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    if not is_voting_open():
        await callback.answer(tr(user_id, "Hozir ovoz berish yopilgan."), show_alert=True)
        return
    if has_voted(user_id):
        await callback.answer(tr(user_id, "Siz faqat 1 marta ovoz bera olasiz."), show_alert=True)
        return
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return
    _, subject_key, teacher_key = parts
    subject_key = normalize_subject_key(subject_key)
    if subject_key not in SUBJECTS or teacher_key not in SUBJECTS[subject_key]["teachers"]:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return

    text = (
        f"❓ <b>Ovozingizni tasdiqlaysizmi?</b>\n\n"
        f"<b>Kafedra:</b> {get_subject_name(subject_key)}\n"
        f"<b>O'qituvchi:</b> {get_teacher_name(subject_key, teacher_key)}\n\n"
        f"Tasdiqlagandan keyin asosiy ovozni qayta o'zgartirib bo'lmaydi."
    )
    await safe_edit_message(callback, tr(user_id, text), confirm_vote_keyboard(user_id, subject_key, teacher_key))
    await callback.answer()


@dp.callback_query(F.data.startswith("confirm_vote:"))
async def confirm_vote_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    if not is_voting_open():
        await callback.answer(tr(user_id, "Hozir ovoz berish yopilgan."), show_alert=True)
        return
    if has_voted(user_id):
        await callback.answer(tr(user_id, "Siz faqat 1 marta ovoz bera olasiz."), show_alert=True)
        return

    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return
    _, subject_key, teacher_key = parts
    subject_key = normalize_subject_key(subject_key)
    if subject_key not in SUBJECTS or teacher_key not in SUBJECTS[subject_key]["teachers"]:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return

    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except TelegramBadRequest:
        pass

    async with db_lock:
        saved = save_vote(
            user_id=user_id,
            full_name=callback.from_user.full_name or "Noma'lum",
            username=callback.from_user.username or "",
            subject_key=subject_key,
            teacher_key=teacher_key
        )

    if not saved:
        await callback.answer(tr(user_id, "Siz allaqachon ovoz bergansiz."), show_alert=True)
        await safe_edit_message(callback, get_already_voted_text(user_id), home_keyboard(user_id))
        return

    text = (
        f"✅ <b>Ovoz muvaffaqiyatli qabul qilindi</b>\n\n"
        f"<b>Bo'lim:</b> {get_subject_name(subject_key)}\n"
        f"<b>Tanlovingiz:</b> {get_teacher_name(subject_key, teacher_key)}\n\n"
        f"Rahmat, sizning ovozingiz saqlandi."
    )
    await safe_edit_message(callback, tr(user_id, text), after_vote_keyboard(user_id))
    await callback.answer(tr(user_id, "Ovozingiz qabul qilindi!"))

# =========================
# RATING CALLBACKS
# =========================
@dp.callback_query(F.data == "go_rating_panel")
async def go_rating_panel_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await safe_edit_message(callback, get_welcome_text(user_id), subscription_keyboard(user_id))
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    await safe_edit_message(callback, get_rating_select_text(user_id), rating_subjects_keyboard(user_id))
    await callback.answer()

@dp.callback_query(F.data.startswith("rating_subject:"))
async def rating_subject_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    subject_key = normalize_subject_key(callback.data.split(":", 1)[1])
    if subject_key not in SUBJECTS:
        await callback.answer(tr(user_id, "Noto'g'ri kafedra."), show_alert=True)
        return
    await safe_edit_message(callback, get_rating_teacher_text(user_id, subject_key), rating_teachers_keyboard(user_id, subject_key))
    await callback.answer()

@dp.callback_query(F.data.startswith("rating_teacher:"))
async def rating_teacher_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return
    _, subject_key, teacher_key = parts
    subject_key = normalize_subject_key(subject_key)
    if subject_key not in SUBJECTS or teacher_key not in SUBJECTS[subject_key]["teachers"]:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return
    await safe_edit_message(callback, get_rate_text(user_id, subject_key, teacher_key), rate_keyboard(user_id, subject_key, teacher_key))
    await callback.answer()

@dp.callback_query(F.data.startswith("rate:"))
async def rate_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not require_access_only(user_id):
        await callback.answer(get_subscription_required_alert(user_id), show_alert=True)
        return
    parts = callback.data.split(":")
    if len(parts) != 4:
        await callback.answer("Noto'g'ri baho.", show_alert=True)
        return
    _, rating, subject_key, teacher_key = parts
    subject_key = normalize_subject_key(subject_key)
    if rating not in ("like", "dislike") or subject_key not in SUBJECTS or teacher_key not in SUBJECTS[subject_key]["teachers"]:
        await callback.answer("Noto'g'ri baho.", show_alert=True)
        return

    async with db_lock:
        save_teacher_rating(user_id, callback.from_user.full_name or "Noma'lum", callback.from_user.username or "", subject_key, teacher_key, rating)

    await safe_edit_message(callback, get_rate_text(user_id, subject_key, teacher_key), rate_keyboard(user_id, subject_key, teacher_key))
    await callback.answer(tr(user_id, "Bahoyingiz saqlandi!"))

# =========================
# USER RESULTS
# =========================
@dp.callback_query(F.data == "show_results_menu_user")
async def show_results_menu_user(callback: CallbackQuery):
    user_id = callback.from_user.id
    await safe_edit_message(callback, get_results_menu_text(user_id, False), results_menu_keyboard_user(user_id))
    await callback.answer()

@dp.callback_query(F.data == "show_results_user:general")
async def show_results_user_general(callback: CallbackQuery):
    user_id = callback.from_user.id
    async with db_lock:
        text = add_refresh_time(get_general_results_text(user_id), user_id)
    await safe_edit_message(callback, text, results_keyboard_user(user_id, "general"))
    await callback.answer()

@dp.callback_query(F.data.startswith("show_results_user:"))
async def show_results_user(callback: CallbackQuery):
    user_id = callback.from_user.id
    scope = normalize_subject_key(callback.data.split(":", 1)[1].strip())

    if scope != "general" and scope not in SUBJECTS:
        await callback.answer(tr(user_id, "Noto'g'ri bo'lim."), show_alert=True)
        return

    async with db_lock:
        text = get_general_results_text(user_id) if scope == "general" else get_subject_results_text(user_id, scope)
        text = add_refresh_time(text, user_id)

    await safe_edit_message(callback, text, results_keyboard_user(user_id, scope))
    await callback.answer()


@dp.callback_query(F.data == "refresh_results_user:general")
async def refresh_results_user_general(callback: CallbackQuery):
    user_id = callback.from_user.id
    refresh_key = "results_user:general"

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        async with db_lock:
            text = add_refresh_time(get_general_results_text(user_id), user_id)
        await safe_edit_message(callback, text, results_keyboard_user(user_id, "general"))
    finally:
        finish_refresh(user_id, refresh_key)

@dp.callback_query(F.data.startswith("refresh_results_user:"))
async def refresh_results_user(callback: CallbackQuery):
    user_id = callback.from_user.id
    scope = normalize_subject_key(callback.data.split(":", 1)[1].strip())
    refresh_key = f"results_user:{scope}"

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)

        if scope != "general" and scope not in SUBJECTS:
            await callback.answer(tr(user_id, "Noto'g'ri bo'lim."), show_alert=True)
            return

        async with db_lock:
            text = get_general_results_text(user_id) if scope == "general" else get_subject_results_text(user_id, scope)
            text = add_refresh_time(text, user_id)

        await safe_edit_message(callback, text, results_keyboard_user(user_id, scope))
    finally:
        finish_refresh(user_id, refresh_key)


# =========================
# ADMIN CALLBACKS
# =========================
@dp.callback_query(F.data == "back_admin_panel")
async def back_admin_panel_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer()

@dp.callback_query(F.data == "admin_results")
async def admin_results_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, get_results_menu_text(user_id, True), results_menu_keyboard_admin(user_id))
    await callback.answer()


@dp.callback_query(F.data == "show_results_admin:general")
async def show_results_admin_general(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    async with db_lock:
        text = get_general_results_text(user_id)
        text = add_refresh_time(text, user_id)
    await safe_edit_message(callback, text, results_keyboard_admin(user_id, "general"))
    await callback.answer()

@dp.callback_query(F.data == "refresh_results_admin:general")
async def refresh_results_admin_general(callback: CallbackQuery):
    user_id = callback.from_user.id
    refresh_key = "results_admin:general"

    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        async with db_lock:
            text = get_general_results_text(user_id)
            text = add_refresh_time(text, user_id)
        await safe_edit_message(callback, text, results_keyboard_admin(user_id, "general"))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data.startswith("show_results_admin:"))
async def show_results_admin(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    scope = normalize_subject_key(callback.data.split(":", 1)[1].strip())
    if scope != "general" and scope not in SUBJECTS:
        await callback.answer(tr(user_id, "Noto'g'ri bo'lim."), show_alert=True)
        return

    async with db_lock:
        text = get_general_results_text(user_id) if scope == "general" else get_subject_results_text(user_id, scope)
        text = add_refresh_time(text, user_id)

    await safe_edit_message(callback, text, results_keyboard_admin(user_id, scope))
    await callback.answer()


@dp.callback_query(F.data.startswith("refresh_results_admin:"))
async def refresh_results_admin_handler(callback: CallbackQuery):
    user_id = callback.from_user.id

    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    scope = normalize_subject_key(callback.data.split(":", 1)[1].strip())
    refresh_key = f"results_admin:{scope}"

    if scope != "general" and scope not in SUBJECTS:
        await callback.answer(tr(user_id, "Noto'g'ri bo'lim."), show_alert=True)
        return

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        async with db_lock:
            text = get_general_results_text(user_id) if scope == "general" else get_subject_results_text(user_id, scope)
            text = add_refresh_time(text, user_id)

        await safe_edit_message(callback, text, results_keyboard_admin(user_id, scope))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data == "admin_rating_stats")
async def admin_rating_stats_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, "⭐️ <b>Baholash foizlari</b>\n\nKerakli bo'limni tanlang:", rating_results_menu_keyboard_admin(user_id))
    await callback.answer()

@dp.callback_query(F.data.startswith("show_rating_stats:"))
async def show_rating_stats_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    scope = normalize_subject_key(callback.data.split(":", 1)[1])
    async with db_lock:
        text = add_refresh_time(get_rating_stats_text(user_id, scope), user_id)
    await safe_edit_message(callback, text, rating_stats_keyboard_admin(user_id, scope))
    await callback.answer()

@dp.callback_query(F.data.startswith("refresh_rating_stats:"))
async def refresh_rating_stats_callback(callback: CallbackQuery):
    user_id = callback.from_user.id

    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    scope = normalize_subject_key(callback.data.split(":", 1)[1])
    refresh_key = f"rating_stats:{scope}"

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        async with db_lock:
            text = add_refresh_time(get_rating_stats_text(user_id, scope), user_id)
        await safe_edit_message(callback, text, rating_stats_keyboard_admin(user_id, scope))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data == "admin_top_ratings")
async def admin_top_ratings_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    kb = InlineKeyboardBuilder()
    kb.row(InlineKeyboardButton(text=tr(user_id, "🔄 Yangilash"), callback_data="admin_top_ratings"))
    kb.row(InlineKeyboardButton(text="⬅️ Admin panel", callback_data="back_admin_panel"))
    await safe_edit_message(callback, get_top_ratings_text(user_id), kb.as_markup())
    await callback.answer()

@dp.callback_query(F.data == "admin_teacher_stats")
async def admin_teacher_stats_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, tr(user_id, "👤 <b>O'qituvchi statistikasi</b>\n\nKafedrani tanlang:"), teacher_stats_subjects_keyboard(user_id))
    await callback.answer()


@dp.callback_query(F.data.startswith("teacher_stats_subject:"))
async def teacher_stats_subject_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    subject_key = normalize_subject_key(callback.data.split(":", 1)[1])
    if subject_key not in SUBJECTS:
        await callback.answer(tr(user_id, "Noto'g'ri kafedra."), show_alert=True)
        return
    await safe_edit_message(
        callback,
        tr(user_id, f"👤 <b>{get_subject_name(subject_key)}</b>\n\nStatistika uchun o'qituvchini tanlang:"),
        teacher_stats_teachers_keyboard(user_id, subject_key)
    )
    await callback.answer()


@dp.callback_query(F.data.startswith("teacher_stats:"))
async def teacher_stats_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer(tr(user_id, "Noto'g'ri tanlov."), show_alert=True)
        return
    _, subject_key, teacher_key = parts
    subject_key = normalize_subject_key(subject_key)
    await safe_edit_message(callback, get_teacher_detailed_stats_text(user_id, subject_key, teacher_key), teacher_stats_keyboard(user_id, subject_key))
    await callback.answer()


@dp.callback_query(F.data == "admin_backup")
async def admin_backup_handler(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await callback.answer("Backup tayyorlanmoqda...")
    try:
        async with db_lock:
            backup_path = create_backup_zip()
        await callback.message.answer_document(
            FSInputFile(backup_path),
            caption="💾 Backup: votes.db, ovozlar Excel, rating Excel va shikoyatlar fayli."
        )
    except Exception as e:
        logging.error(f"Backup yaratishda xatolik: {e}")
        await callback.message.answer(f"❌ Backup yaratishda xatolik: {e}")


@dp.callback_query(F.data == "admin_complaints")
async def admin_complaints_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    async with db_lock:
        text = add_refresh_time(get_complaints_text(user_id), user_id)
    await safe_edit_message(callback, text, complaints_keyboard_admin(user_id))
    await callback.answer()


@dp.callback_query(F.data == "refresh_admin_complaints")
async def refresh_admin_complaints_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    refresh_key = "admin_complaints"

    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        async with db_lock:
            text = add_refresh_time(get_complaints_text(user_id), user_id)
        await safe_edit_message(callback, text, complaints_keyboard_admin(user_id))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data == "admin_export_complaints_docx")
async def admin_export_complaints_docx_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    filename = export_complaints_to_docx()
    await callback.message.answer_document(
        FSInputFile(filename),
        caption="📄 Shikoyat va takliflar Word fayl ko'rinishida."
    )
    await callback.answer()


@dp.callback_query(F.data == "admin_users")
async def admin_users_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    async with db_lock:
        text = add_refresh_time(get_users_text(user_id), user_id)
    await safe_edit_message(callback, text, users_keyboard_admin(user_id))
    await callback.answer()

@dp.callback_query(F.data == "refresh_admin_users")
async def refresh_admin_users(callback: CallbackQuery):
    user_id = callback.from_user.id
    refresh_key = "admin_users"

    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return

    if not can_start_refresh(user_id, refresh_key):
        await callback.answer(tr(user_id, "Juda tez bosyapsiz. 1-2 soniyadan keyin urinib ko'ring."), show_alert=False)
        return

    try:
        await callback.answer(tr(user_id, "Yangilanmoqda..."), show_alert=False)
        async with db_lock:
            text = add_refresh_time(get_users_text(user_id), user_id)
        await safe_edit_message(callback, text, users_keyboard_admin(user_id))
    finally:
        finish_refresh(user_id, refresh_key)


@dp.callback_query(F.data == "admin_export_votes_excel")
async def admin_export_votes_excel_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    filename = export_votes_to_excel()
    await callback.message.answer_document(FSInputFile(filename), caption="📁 Ovozlar Excel fayl ko'rinishida.")
    await callback.answer()

@dp.callback_query(F.data == "admin_export_rating_excel")
async def admin_export_rating_excel_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    filename = export_rating_to_excel()
    await callback.message.answer_document(FSInputFile(filename), caption="📁 Rating Excel fayl ko'rinishida.")
    await callback.answer()

@dp.callback_query(F.data == "admin_open")
async def admin_open_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    open_voting()
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer("Voting ochildi!")

@dp.callback_query(F.data == "admin_close")
async def admin_close_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    close_voting()
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer("Voting yopildi!")

@dp.callback_query(F.data == "admin_reset_votes_confirm")
async def admin_reset_votes_confirm_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, "⚠️ <b>Diqqat!</b>\n\nBarcha ovozlar o'chiriladi.\nDavom etasizmi?", reset_confirm_keyboard(user_id, "votes"))
    await callback.answer()

@dp.callback_query(F.data == "admin_reset_rating_confirm")
async def admin_reset_rating_confirm_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, "⚠️ <b>Diqqat!</b>\n\nBarcha rating baholari o'chiriladi.\nDavom etasizmi?", reset_confirm_keyboard(user_id, "rating"))
    await callback.answer()

@dp.callback_query(F.data == "admin_reset_complaints_confirm")
async def admin_reset_complaints_confirm_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(
        callback,
        "⚠️ <b>Diqqat!</b>\n\nBarcha shikoyat va takliflar o'chiriladi.\nDavom etasizmi?",
        reset_confirm_keyboard(user_id, "complaints")
    )
    await callback.answer()

@dp.callback_query(F.data == "admin_reset_complaints")
async def admin_reset_complaints_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    async with db_lock:
        reset_complaints()
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer("Shikoyat va takliflar tozalandi!")

@dp.callback_query(F.data == "cancel_reset")
async def cancel_reset_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer(tr(user_id, "Bekor qilindi"))

@dp.callback_query(F.data == "admin_reset_votes")
async def admin_reset_votes_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    async with db_lock:
        reset_votes()
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer("Ovozlar reset qilindi!")

@dp.callback_query(F.data == "admin_reset_rating")
async def admin_reset_rating_callback(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not is_admin(user_id):
        await callback.answer("Siz admin emassiz.", show_alert=True)
        return
    async with db_lock:
        reset_ratings()
    await safe_edit_message(callback, get_admin_panel_text(user_id), admin_panel_keyboard(user_id))
    await callback.answer("Rating reset qilindi!")


# FIX #6: Keraksiz takroriy callbacklar olib tashlandi.
# admin_complaints_word, admin_export_complaints_word, admin_export, admin_reset_confirm,
# admin_reset — bular yo'q yoki birining nomi bilan almashtirilib qolgan edi.
# Eski nomlar hali ham ishlashi kerak bo'lsa, quyida saqlanadi:

@dp.callback_query(F.data == "admin_export")
async def admin_export_callback(callback: CallbackQuery):
    await admin_export_votes_excel_callback(callback)

@dp.callback_query(F.data == "admin_reset_confirm")
async def admin_reset_confirm_old_callback(callback: CallbackQuery):
    await admin_reset_votes_confirm_callback(callback)

@dp.callback_query(F.data == "admin_reset")
async def admin_reset_old_callback(callback: CallbackQuery):
    await admin_reset_votes_callback(callback)

# =========================
# TEXT HANDLER
# =========================
@dp.message(F.text)
async def text_handler(message: Message):
    user_id = message.from_user.id
    ensure_user(user_id)

    if user_id in WAITING_COMPLAINT_TEXT:
        text = (message.text or "").strip()
        if not text:
            await message.answer("Bo'sh xabar qabul qilinmaydi. Iltimos, matn yozing.")
            return

        allowed, reason = complaint_allowed(user_id, text)
        if not allowed:
            await message.answer(tr(user_id, reason), parse_mode="HTML", reply_markup=complaint_cancel_keyboard(user_id))
            return

        WAITING_COMPLAINT_TEXT.discard(user_id)
        async with db_lock:
            save_complaint(
                user_id=user_id,
                full_name=message.from_user.full_name or "Noma'lum",
                username=message.from_user.username or "",
                message_text=text,
            )
        await message.answer(
            get_complaint_saved_text(user_id),
            parse_mode="HTML",
            reply_markup=home_keyboard(user_id)
        )
        return

    if message.text and message.text.lower() == "results":
        await message.answer(
            get_results_menu_text(user_id, False),
            parse_mode="HTML",
            reply_markup=results_menu_keyboard_user(user_id)
        )


# =========================
# MAIN
# =========================
async def main():
    init_db()
    logging.info(f"Bot ishga tushdi. Baza: {DB_NAME}")
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
